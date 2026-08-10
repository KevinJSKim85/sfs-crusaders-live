#!/usr/bin/env python3
"""Builds the operator manual .docx from the content below.

Run after any change that alters how a panel gets its data:
    python3 docs/build_manual.py
Output: ../SFS_Crusader_Hub_Operator_Manual.docx (next to the repo folder).
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEAD_FONT = "Avenir Next"
BODY_FONT = "Helvetica Neue"
RED = RGBColor(0xA6, 0x19, 0x2E)
INK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x5A, 0x5A, 0x5A)

UPDATED = "10 August 2026"

doc = Document()

# ---------- base styles ----------
for name, font, size in (("Normal", BODY_FONT, 10.5),
                         ("List Paragraph", BODY_FONT, 10.5)):
    st = doc.styles[name]
    st.font.name = font
    st.font.size = Pt(size)
    st.font.color.rgb = INK
    st.element.rPr.rFonts.set(qn("w:eastAsia"), font)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.12

for name, size, color, before in (("Heading 1", 15, RED, 18),
                                  ("Heading 2", 11.5, INK, 12)):
    st = doc.styles[name]
    st.font.name = HEAD_FONT
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.element.rPr.rFonts.set(qn("w:eastAsia"), HEAD_FONT)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(5)
    st.paragraph_format.keep_with_next = True

for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.8)
    s.left_margin = s.right_margin = Inches(0.85)


def para(text="", style=None, size=None, bold=False, color=None,
         space_after=None, font=BODY_FONT):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.name = font
    r.element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def bullets(items):
    for it in items:
        para(it, style="List Bullet")


def h1(t):
    doc.add_paragraph(t, style="Heading 1")


def h2(t):
    doc.add_paragraph(t, style="Heading 2")


def mono(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.28)
    r = p.add_run(text)
    r.font.name = "Menlo"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x30, 0x30, 0x30)
    return p


def shade(cell, hexfill):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hexfill)
    cell._tc.get_or_add_tcPr().append(el)


def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(htext)
        r.font.name = HEAD_FONT
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, "A6192E")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(val)
            r.font.name = BODY_FONT
            r.font.size = Pt(9)
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)
        for c in row.cells:
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
            c.paragraphs[0].paragraph_format.space_before = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ---------- cover block ----------
para("SFS CRUSADER HUB", size=21, bold=True, color=RED, space_after=0,
     font=HEAD_FONT)
para("Operator Manual", size=13, bold=True, space_after=1, font=HEAD_FONT)
para("How each panel gets its data, and what you have to do by hand",
     size=10.5, color=GREY, space_after=9)
para("Site:  sfs-crusaders-live.vercel.app", size=9.5, color=GREY, space_after=1)
para("Admin:  sfs-crusaders-live.vercel.app/admin  (Google sign-in, allowlisted "
     "accounts only)", size=9.5, color=GREY, space_after=1)
para("Code:  github.com/KevinJSKim85/sfs-crusaders-live", size=9.5, color=GREY,
     space_after=1)
para("Updated:  " + UPDATED, size=9.5, color=GREY, space_after=6)

# ---------- 1 ----------
h1("1. Read this first")
bullets([
    "Everything on the dashboard now updates itself, including the lunch menu. "
    "There is no weekly task.",
    "The only thing you post by hand is a club announcement, and only when you "
    "have one. Optionally you can paste the weekly HS newsletter link.",
    "Every panel shows when it was last updated, in its top-right corner. If a "
    "panel looks wrong, read that label before anything else.",
    "You cannot break the site by leaving it alone. Panels keep showing their "
    "last good content and say how old it is.",
    "You never need to redeploy or touch code to change what students see in "
    "the announcements or newsletter slots. Both are edited in the admin page.",
])

# ---------- 2 ----------
h1("2. How refreshing works")
para("The same rules apply to every panel, so you only have to learn them once.")

h2("Two stages")
bullets([
    "Server side. The site fetches each source once and holds the result for a "
    "set time. Every visitor shares that one copy, so a thousand students can "
    "open the page and the source is still asked only once.",
    "Browser side. Each open tab asks the site for fresh data on its own timer. "
    "That timer is in the table below.",
])
para("Result: a change at the source can take up to the server hold time plus "
     "the browser timer to appear. For most panels that is under half an hour.")

h2("Behaviour worth knowing")
bullets([
    "A tab left open in the background stops refreshing. It catches up the "
    "moment you switch back to it.",
    "If a source is down, the panel keeps the content it already had and "
    "retries less and less often, up to 30 minutes apart. It does not blank out.",
    "Panels paint instantly from the last copy stored in that browser, then "
    "update in place. Reopening the site is never a blank wait.",
    "Reloading the page does not force fresh data from the source. The server "
    "copy is still within its hold time. This is normal.",
])

# ---------- 3 ----------
h1("3. Every panel at a glance")
table(
    ["Panel", "Where the data comes from", "Refreshes", "Who updates it"],
    [
        ["High School Lunch", "seoulforeign.org/lunch, High School PDF",
         "Every 30 min", "Nobody"],
        ["Upcoming Calendar", "seoulforeign.org school calendar",
         "Every 6 hours", "Nobody"],
        ["Clubs & Announcements", "Admin posts; falls back to the SFHS clubs site",
         "Instant / 6 hours", "You, as needed"],
        ["SFS News", "seoulforeign.org news page, plus the newsletter link you paste",
         "Every 15 min", "The school / you, optional"],
        ["Athletics Results", "sfscrusaderathletics.com", "Every 15 min",
         "Athletics dept."],
        ["Instagram (x2)", "The two Instagram accounts", "Every 30 min",
         "Whoever posts"],
        ["World & Reading", "Student's chosen sources, out of 18",
         "Every 30 min", "Each student"],
        ["College News", "Student's chosen sources, out of 18", "Every 30 min",
         "Each student"],
        ["Spirit Magazine", "issuu.com/thesfhsspirit", "Every hour",
         "The Spirit staff"],
        ["Weather & Air Quality", "Open-Meteo (in the top banner)",
         "Every 10 min", "Nobody"],
        ["Notes", "Saved in the student's own browser", "As they type",
         "Each student"],
        ["Quick Links", "Fixed list in the code", "Never", "Developer"],
    ],
    [1.35, 2.75, 1.05, 1.55],
)

# ---------- 4 ----------
h1("4. What you have to do")

h2("As needed: post an announcement")
bullets([
    "In /admin, fill in the club name, the message, and an optional link. Post.",
    "It appears on the dashboard immediately. No waiting, no reload. This panel "
    "is live.",
    "When there are no announcements, the panel shows the club directory "
    "instead, so it is never empty.",
    "Keep messages short. Anything past roughly 110 characters is trimmed in "
    "the panel.",
])

h2("Optional, weekly: link the HS newsletter")
bullets([
    "The newsletter is emailed, not published to a public archive, so nothing "
    "can fetch it automatically. That is a limitation of how it is sent, not of "
    "the site.",
    "If you want it on the dashboard: copy the 'View this email in your browser' "
    "link out of the email, paste it into the Newsletter field in /admin with a "
    "title, and save.",
    "It shows as the first card in SFS News, marked Newsletter. Paste a new link "
    "each week; the old one is replaced.",
    "Leaving it blank is fine. The panel simply shows school news only.",
])

h2("Rarely: after changing database permissions")
bullets([
    "Editing firestore.rules and pushing to GitHub does nothing on its own. The "
    "file has to be deployed:",
])
mono("firebase deploy --only firestore:rules")
bullets([
    "Skipping this leaves the old permissions live while the file looks "
    "correct. Do not skip it.",
])

h2("No longer your job: the lunch menu")
bullets([
    "The site reads the High School menu straight from seoulforeign.org/lunch. "
    "When the school posts a new PDF there, the dashboard picks it up within 30 "
    "minutes, on any day, at any hour.",
    "There is still a Drive link field in /admin. It is an override, not the "
    "normal path. Use it only if the school page breaks, or to publish a menu "
    "before they post one.",
    "If you do fill it in, it wins over the school page until you clear it. "
    "Clear it when you are done.",
])

# ---------- 5 ----------
h1("5. Reading the status labels")
para("Each panel shows one of these in its top-right corner.")
table(
    ["Label", "Meaning", "What to do"],
    [
        ["Updated 3m ago", "Working normally.", "Nothing"],
        ["Updated 2h ago · retrying",
         "The source did not answer. You are seeing the last good copy.",
         "Wait. It retries on its own"],
        ["Offline · 14:32", "The device has no internet.", "Check the network"],
        ["Unavailable",
         "The source failed and there was no earlier copy to fall back on.",
         "See section 7"],
        ["Season break",
         "On Athletics: the newest result is over a month old.",
         "Nothing, it is out of season"],
        ["not updated recently",
         "On Lunch: the school has not posted a new PDF in over 10 days.",
         "Nothing. Check the school page if it persists"],
    ],
    [1.5, 3.0, 2.2],
)

# ---------- 6 ----------
h1("6. Panel by panel")

h2("High School Lunch")
bullets([
    "Opens seoulforeign.org/lunch, finds the link labelled High School, follows "
    "it to whatever PDF is currently behind it, and reads the menu out of it.",
    "The link is found by its label, not by a saved address, so the school can "
    "replace the file as often as they like and nothing here has to change.",
    "Shows only today's items, by Seoul time, not the device's time zone.",
    "On weekends it says lunch resumes Monday instead of listing weekly items.",
    "The date in the PDF's filename is often wrong. The panel ignores it and "
    "uses the school's upload time instead.",
    "Server holds the parsed result for 30 minutes; browsers ask every 30 minutes.",
    "If the school page cannot be reached, it falls back to the Drive link in "
    "/admin, if one is set.",
])

h2("Upcoming Calendar")
bullets([
    "Reads the school calendar for this month and next, and lists the next "
    "seven things.",
    "Each entry is tagged Arts, Academic, No school, Athletics or School.",
    "Athletics is hidden by default, because the Athletics panel already covers "
    "it. The small button in the panel's corner turns it on for that student.",
    "Multi-day things, like a three-day audition block, are shown as one entry "
    "with a date range, not as three repeats.",
    "The calendar also publishes day-cycle codes (1, 1A, 6A). Those are "
    "filtered out.",
    "Nothing to maintain. It follows the school calendar.",
])

h2("Clubs & Announcements")
bullets([
    "Announcements post instantly. This panel uses a live connection, not a "
    "timer.",
    "With no announcements, it lists all 71 clubs grouped by Day 1, Day 4, "
    "Official Functions, Honor Societies, Pursuits and Academies, plus the "
    "sign-up form and club calendar.",
    "The club list is read from the SFHS clubs Google Site and refreshes every "
    "6 hours. Add or remove a club there and it appears here on its own.",
    "Club names are read from the page structure, not typed into the code, so "
    "next year's roster needs no code change.",
])

h2("SFS News")
bullets([
    "Reads the school's news page and sorts by each article's real publication "
    "date.",
    "If you have pasted a newsletter link, that card sits at the top, marked "
    "Newsletter.",
    "Worth knowing: the school has published twice in the last eighteen months. "
    "The panel is showing everything there is. If it looks thin, that is the "
    "source, not a fault.",
])

h2("Athletics Results")
bullets([
    "Reads the scoreboard data from the athletics site: real scores and "
    "opponents, not a summary page.",
    "When the newest result is more than 30 days old, a Season break line "
    "appears above it so old scores are not mistaken for this week's.",
])

h2("Instagram (two panels)")
bullets([
    "Reads the two accounts directly. No password, no token, nothing that "
    "expires.",
    "Shows the eight most recent posts. Photos and reels both work.",
    "The school account has a backup source if the main one fails; it can lag "
    "by a post or two.",
    "Posting to Instagram is all that is needed. The panel picks it up within "
    "30 minutes.",
])

h2("World & Reading, College News")
bullets([
    "18 sources are available, including four New York Times sections, BBC, "
    "Guardian, Economist, Quanta, NASA, Korea Herald, Inside Higher Ed, "
    "Hechinger, Harvard Gazette and MIT News.",
    "Each student picks their own with the small button in the panel's corner. "
    "The choice is saved on that device only and does not affect anyone else.",
    "Out of the box, World & Reading shows world news and College News shows "
    "education news.",
    "If one source is down the others still show. No action needed, ever.",
])

h2("Spirit Magazine")
bullets([
    "Reads the magazine's Issuu page. A new issue appears within the hour of "
    "being published there.",
    "Each cover is matched to its issue from the same place on the page, so the "
    "cover and title cannot drift apart.",
])

h2("Weather & Air Quality")
bullets([
    "Seoul weather and air quality, refreshed every 10 minutes. Nothing to "
    "maintain.",
    "When air quality data is missing it shows Unavailable rather than a "
    "number. A wrong reading is worse than none.",
])

h2("Notes")
bullets([
    "Saved in each student's own browser. Not shared, not backed up, not "
    "visible to you.",
    "Clearing browser data deletes it. Worth telling students once.",
])

h2("Quick Links")
bullets([
    "Ordered by how often a high school student actually needs them: Gmail, "
    "ManageBac, Classroom, Drive, Calendar, Clubs, Maia Learning, NoodleTools, "
    "Library, Sora, Padlet, then the rest.",
    "Changing the list or the order is a code change. Ask a developer.",
])

h2("Panel layout")
bullets([
    "On a computer, students can drag a panel by its title bar to rearrange "
    "their own dashboard. The arrangement is saved on that device only.",
    "The circular arrow in the top bar resets it. Dragging is off on phones and "
    "tablets, so scrolling never moves panels by accident.",
])

# ---------- 7 ----------
h1("7. When something looks wrong")

h2("One panel says Unavailable")
bullets([
    "Check whether the source itself is up by opening it directly: the "
    "athletics site, the Issuu page, and so on.",
    "To see the exact reason, open the matching address, e.g. /api/athletics. "
    "It returns a short report including an error field.",
    "If the source is up and the panel is not, it needs a developer. Send them "
    "that error field.",
])

h2("Every panel says Unavailable")
bullets([
    "Usually the hosting is down. Check status.vercel.com.",
    "Weather is fetched separately, so if weather works and everything else "
    "fails, the problem is the site rather than the network.",
])

h2("An announcement will not post")
bullets([
    "Confirm you are signed in with an allowlisted account. Others can reach "
    "the page but cannot save.",
    "Adding an operator means changing the allowlist in two places and "
    "redeploying the permissions file. That is a developer task.",
])

h2("The lunch menu is blank or shows last week")
bullets([
    "First open seoulforeign.org/lunch yourself and check the High School link "
    "actually leads to the current menu. Most of the time the school has not "
    "posted yet, and there is nothing to fix.",
    "Open /api/lunch. If it says origin: school, the site is reading the school "
    "page correctly. If it says origin: admin, someone left an override in the "
    "Drive field; clear it.",
    "If the school changed the PDF layout heavily, the reader may not recognise "
    "the sections. The panel falls back to showing the PDF itself, which is "
    "still usable while it gets fixed.",
    "As a stopgap you can always paste a Drive link in /admin. It takes over "
    "immediately.",
])

# ---------- 8 ----------
h1("8. Known gaps")
bullets([
    "Assessment calendars for Grades 9-12 are not linked. The four links are "
    "not published anywhere findable. Send them to a developer and they can be "
    "added.",
    "The HS newsletter cannot be fetched automatically because it only exists "
    "as an email. Pasting the link is the only way to get it on the dashboard.",
    "The Spirit panel keeps a small built-in list of recent issues for "
    "first-time visitors. It is refreshed from Issuu straight away, so it only "
    "matters for the first second of a first visit.",
    "Announcements have never been used. The panel works; nothing has been "
    "posted to it.",
])

# ---------- 9 ----------
h1("9. If you remember nothing else")
bullets([
    "There is no recurring task any more. Lunch updates itself from the school "
    "page.",
    "Post announcements in /admin when you have one. They appear instantly.",
    "Optionally paste the newsletter link each week. Skipping it breaks "
    "nothing.",
    "Check the label in a panel's corner before reporting a problem.",
    "If you edit firestore.rules, deploy it, or the change is not live.",
])

out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "..", "..", "SFS_Crusader_Hub_Operator_Manual.docx")
out = os.path.normpath(out)
doc.save(out)
print("wrote", out)
